from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from openpyxl import load_workbook
import win32com.client as win32

outlook = win32.Dispatch("outlook.application")

#Abre Arquivo Excel
nomeArquivExcel = "DadosAlunosEmail.xlsx"
planilhaDadosAlunos = load_workbook(nomeArquivExcel)

#Selecionando a aba
sheetSelecionada = planilhaDadosAlunos["Nomes"]

for linha in range(2,len(sheetSelecionada["A"]) + 1):

    #Abre Arquivo Word
    arquivoWord = Document("Certificado3.docx")

    #Seleciona o estilo
    estilo = arquivoWord.styles["Normal"]

    #Recuperar nome aluno
    nomeAluno = sheetSelecionada['A%s' % linha].value
    dia = sheetSelecionada['B%s' % linha].value
    mes = sheetSelecionada['C%s' % linha].value
    ano = sheetSelecionada['D%s' % linha].value
    nomeCurso = sheetSelecionada['E%s' % linha].value
    nomeInstrutor = sheetSelecionada['F%s' % linha].value
    emailAluno = sheetSelecionada['G%s' % linha].value

    for paragrafo in arquivoWord.paragraphs:

        if "@nome" in paragrafo.text:
            paragrafo.text = nomeAluno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        paragrafoP1 = "Concluiu com sucesso o curso de "
        paragrafoP2 = ", como carga horária de 20 horas, promovido pela escola de Cursos Online em "
        terceiraParteParagrafo = paragrafoP2 + str(dia) + " de " + mes + " de " + str(ano) + "."

        if "escola" in paragrafo.text:
            paragrafo.text = paragrafoP1
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)
            adicionaNovaPalavra = paragrafo.add_run(nomeCurso)
            #Mudar cor para vermelho
            adicionaNovaPalavra.font.color.rgb = RGBColor(255, 0, 0)
            #Sublinhado
            adicionaNovaPalavra.underline = True
            #Negrito
            adicionaNovaPalavra.bold = True
            adicionaNovaPalavra = paragrafo.add_run(terceiraParteParagrafo)
            #Mudar para cor preto
            adicionaNovaPalavra.font.color.rgb = RGBColor(0, 0, 0)

        if "Instrutor" in paragrafo.text:
            paragrafo.text = nomeInstrutor + " - Instrutor"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    #Recuperar caminho pasta certificados
    caminhoCertificados = "C:\\Users\\lucas\\Documents\\Curso RPA Python\\pythonProject\\Word\\Certificados\\" + nomeAluno + ".docx"

    #Salvando o arquivo
    arquivoWord.save(caminhoCertificados)

    #Recuperar primeiro nome do aluno
    primeiroNome = nomeAluno.split(None,1)[0]

    emailOutlook = outlook.CreateItem(0)
    emailOutlook.To = emailAluno
    emailOutlook.Subject = "Certificado " + nomeAluno
    emailOutlook.HTMLBody = f"""
        <p>Boa noite {primeiroNome}.</p>
        <p>Segue seu <b>certificado</b></p>
        <p>Atenciosamente,</p>
    """

    #Adicionar anexo
    emailOutlook.Attachments.Add(caminhoCertificados)

    #.save() = Salvar como Draft no Outlook
    #.send ou .send() = Enviar o email
    emailOutlook.save()

print("Certificados gerados e enviados com sucesso")